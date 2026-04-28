"""Build the Spindle user manual DOCX from the canonical Markdown manual."""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "spindle_user_manual.md"
OUT = ROOT / "docs" / f"Spindle工业嵌入式Agent使用手册_{date.today().isoformat()}.docx"

FONT_CN = "宋体"
FONT_HEADING = "黑体"
FONT_CODE = "Consolas"


def set_run_font(run, *, name: str = FONT_CN, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 6, line: float = 1.18) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    set_paragraph_spacing(paragraph, after=0, line=1.12)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_grid(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "left", "bottom", "right"):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "120")
                node.set(qn("w:type"), "dxa")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=8 if level == 1 else 5, after=5)
    run = paragraph.add_run(text)
    size = 15 if level == 1 else 12.5
    set_run_font(run, name=FONT_HEADING, size=size, bold=True)
    paragraph.style = f"Heading {min(level, 3)}"
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=6)
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    add_inline_runs(paragraph, text)


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name=FONT_CODE, size=9.5)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=10.5)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_grid(table)
    cell = table.cell(0, 0)
    shade_cell(cell, "F2F3F5")
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=0, line=1.08)
    for i, line in enumerate(lines):
        if i:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_run_font(run, name=FONT_CODE, size=8.7)
    doc.add_paragraph()


def add_bullet(doc: Document, text: str, ordered: bool = False) -> None:
    paragraph = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
    set_paragraph_spacing(paragraph, after=3)
    add_inline_runs(paragraph, text)


def add_note_box(doc: Document, title: str, lines: list[str], fill: str = "EEF5EF") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_grid(table)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=2, line=1.14)
    run = paragraph.add_run(title)
    set_run_font(run, name=FONT_HEADING, size=10.5, bold=True)
    for line in lines:
        p = cell.add_paragraph()
        set_paragraph_spacing(p, after=1, line=1.12)
        add_inline_runs(p, line)
    doc.add_paragraph()


def add_info_table(doc: Document) -> None:
    rows = [
        ("文档名称", "Spindle 工业嵌入式 Agent 使用手册"),
        ("版本日期", date.today().isoformat()),
        ("适用范围", "Spindle v1.0 稳定版；资料检索、风险审计、离线回归、开发板只读诊断"),
        ("硬件边界", "默认不执行真实硬件动作；涉及 --execute、总线、运动、IO、remoteproc、固件时必须进入硬件窗口"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_grid(table)
    for row_idx, (key, value) in enumerate(rows):
        set_cell_text(table.cell(row_idx, 0), key, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(row_idx, 1), value)
        shade_cell(table.cell(row_idx, 0), "F0F0F0")
    doc.add_paragraph()


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    for style_name, size in (("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11.5)):
        style = styles[style_name]
        style.font.name = FONT_HEADING
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.55)
    section.right_margin = Cm(2.35)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Spindle 工业嵌入式 Agent 使用手册")
    set_run_font(run, size=9)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Spindle\n工业嵌入式 Agent 使用手册")
    set_run_font(run, name=FONT_HEADING, size=24, bold=True)
    set_paragraph_spacing(title, after=12, line=1.15)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("资料检索、风险审计、离线回归与开发板只读诊断")
    set_run_font(run, size=12)
    set_paragraph_spacing(subtitle, after=18)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"版本日期：{date.today().isoformat()}")
    set_run_font(run, size=10.5)
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_front_matter(doc: Document) -> None:
    add_heading(doc, "文档说明", 1)
    add_info_table(doc)
    add_note_box(
        doc,
        "使用边界",
        [
            "Spindle 用于把真实 EtherCAT、机器人、RPMsg、M7、fake harness 经验沉淀为可检索、可回归、可审计的工程工具。",
            "Spindle 不替代真实上板验证，也不自动执行机器人运动、IO 输出、remoteproc 生命周期操作或固件热重载。",
        ],
        fill="F7F7F7",
    )


def build_from_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    in_code = False
    code_lines: list[str] = []
    pending_para: list[str] = []

    def flush_para() -> None:
        nonlocal pending_para
        if pending_para:
            add_body_paragraph(doc, " ".join(pending_para).strip())
            pending_para = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                flush_para()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            flush_para()
            continue
        if line.startswith("# "):
            flush_para()
            continue
        if line.startswith("## "):
            flush_para()
            add_heading(doc, line[3:].strip(), 1)
            continue
        if line.startswith("### "):
            flush_para()
            add_heading(doc, line[4:].strip(), 2)
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            flush_para()
            add_bullet(doc, numbered.group(1), ordered=True)
            continue
        if line.startswith("- "):
            flush_para()
            add_bullet(doc, line[2:].strip())
            continue
        pending_para.append(line.strip())
    flush_para()


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    setup_styles(doc)
    setup_document(doc)
    add_cover(doc)
    add_front_matter(doc)
    build_from_markdown(doc, markdown)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
